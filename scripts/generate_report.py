"""
Generate the Engineering Improvements Report as a Word (.docx) document.
Run: python3 scripts/generate_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Palette ────────────────────────────────────────────────────────────────────
INDIGO   = RGBColor(0x63, 0x66, 0xF1)   # #6366F1  accent
DARK     = RGBColor(0x1E, 0x1E, 0x2E)   # near-black heading
SLATE    = RGBColor(0x47, 0x55, 0x69)   # muted body
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = RGBColor(0xF1, 0xF5, 0xF9)  # table zebra
GREEN    = RGBColor(0x16, 0xA3, 0x4A)
RED      = RGBColor(0xDC, 0x26, 0x26)

# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def add_heading(doc, text, level=1, color=DARK, size=None, bold=True, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    default_sizes = {1: 22, 2: 16, 3: 13}
    run.font.size = Pt(size or default_sizes.get(level, 12))
    return p


def add_body(doc, text, color=SLATE, size=10.5, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text, color=SLATE, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 90)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    run.font.size = Pt(7)


def add_colored_heading(doc, text, bg_hex="6366F1", text_color=WHITE):
    """Full-width shaded paragraph acting as a section banner."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_hex)
    pPr.append(shd)
    # Padding via spacing
    pPr_spacing = OxmlElement("w:spacing")
    pPr_spacing.set(qn("w:before"), "80")
    pPr_spacing.set(qn("w:after"), "80")
    pPr.append(pPr_spacing)
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.color.rgb = text_color
    run.font.size = Pt(12)
    return p


# ── Document ───────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover / Title ──────────────────────────────────────────────────────────────
doc.add_paragraph()  # top pad

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Agentic S3 Website Builder")
title_run.bold = True
title_run.font.size = Pt(28)
title_run.font.color.rgb = INDIGO

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle_p.add_run("Engineering Improvements Report")
sub_run.font.size = Pt(16)
sub_run.font.color.rgb = DARK

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_p.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}")
date_run.font.size = Pt(10)
date_run.font.color.rgb = SLATE
date_run.italic = True

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

# ── Executive Summary ──────────────────────────────────────────────────────────
add_heading(doc, "Executive Summary", level=1, color=INDIGO, size=18, space_before=6)
add_body(doc, (
    "Between March and April 2026, the Agentic S3 Website Builder underwent four "
    "structured rounds of engineering improvements — labelled Priority 1 through "
    "Priority 4. Each priority addressed a distinct dimension of product quality: "
    "Security, Stability, Code Quality, and Performance & User Experience. "
    "This document explains what was done in each phase, why the work was undertaken, "
    "and the measurable and qualitative benefits delivered."
))

doc.add_paragraph()

# ── High-level summary table ───────────────────────────────────────────────────
add_heading(doc, "Phase Overview", level=2, color=DARK, size=13, space_before=10)

headers = ["Priority", "Theme", "Commits", "Key Outcome"]
rows = [
    ["P1", "Security Hardening", "07bae09", "Eliminated OWASP Top-10 vulnerabilities"],
    ["P2", "Stability & Reliability", "06310e3", "Resilient uploads, retries, health checks"],
    ["P3", "Code Quality & Testing", "d3a0aba", "22-test pytest suite, modular frontend JS"],
    ["P4", "Performance & UX Polish", "153b550", "Faster editor, real-time build stream, toast UI"],
]

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"

# Header row
hdr_cells = tbl.rows[0].cells
for i, h in enumerate(headers):
    _set_cell_bg(hdr_cells[i], "6366F1")
    _cell_margins(hdr_cells[i])
    p = hdr_cells[i].paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

# Data rows
for ri, row_data in enumerate(rows):
    cells = tbl.add_row().cells
    bg = "F1F5F9" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        _set_cell_bg(cells[ci], bg)
        _cell_margins(cells[ci])
        p = cells[ci].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9.5)
        if ci == 0:
            run.bold = True
            run.font.color.rgb = INDIGO

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# PRIORITY 1 — SECURITY
# ══════════════════════════════════════════════════════════════════════════════
add_colored_heading(doc, "Priority 1 — Security Hardening  (commit 07bae09)")

add_heading(doc, "What Was Done", level=2, color=DARK, size=12, space_before=10)
bullets_p1_what = [
    "Replaced all raw SQL string concatenation with parameterised queries throughout the database layer.",
    "Added input-length and format validation on every API boundary (email, URLs, passwords, IDs).",
    "Enforced bcrypt password hashing with a minimum cost factor of 12 everywhere credentials are stored.",
    "Introduced HTTPS-only, HttpOnly, SameSite=Strict cookie flags for session tokens.",
    "Removed hard-coded secrets and API keys from source code; migrated them to environment variables.",
    "Applied rate-limiting middleware to authentication endpoints (login, OTP, password reset).",
    "Tightened CORS policy: restricted allowed origins to explicitly configured domains.",
    "Added Content-Security-Policy and X-Frame-Options response headers.",
]
for b in bullets_p1_what:
    add_bullet(doc, b)

add_heading(doc, "Why It Was Done", level=2, color=DARK, size=12, space_before=8)
add_body(doc, (
    "Several OWASP Top-10 2021 categories were directly applicable to the codebase: "
    "A03 (Injection) via unsanitised SQL, A02 (Cryptographic Failures) via weak hashing, "
    "A07 (Identification & Authentication Failures) via absent rate-limiting, and "
    "A05 (Security Misconfiguration) via permissive CORS and missing security headers. "
    "Addressing these before adding new features is standard secure-development practice — "
    "retrofitting security is significantly more expensive than building it in from the start."
))

add_heading(doc, "Benefits", level=2, color=DARK, size=12, space_before=8)
benefits_p1 = [
    "Eliminated SQL-injection attack surface across all database operations.",
    "Password breaches now yield only bcrypt hashes — brute-force cost increases exponentially.",
    "Rate-limiting prevents credential stuffing and brute-force login attacks.",
    "Proper CORS and CSP headers protect users against XSS and clickjacking.",
    "No secrets in version control — rotating keys no longer requires a code deployment.",
]
for b in benefits_p1:
    add_bullet(doc, b)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PRIORITY 2 — STABILITY
# ══════════════════════════════════════════════════════════════════════════════
add_colored_heading(doc, "Priority 2 — Stability & Reliability  (commit 06310e3)", bg_hex="0F766E")

add_heading(doc, "What Was Done", level=2, color=DARK, size=12, space_before=10)
bullets_p2_what = [
    "Wrapped all S3 upload calls in exponential-backoff retry logic (3 attempts, jitter).",
    "Added a database migration versioning table so schema changes are tracked and idempotent.",
    "Introduced OpenAI API retry with back-off to handle transient rate-limit (429) responses.",
    "Implemented a GET /health endpoint returning service status, database connectivity, and version.",
    "Added graceful startup and shutdown hooks to cleanly close DB pools and background tasks.",
    "Improved error propagation: unhandled exceptions now log a full stack trace instead of swallowing errors.",
]
for b in bullets_p2_what:
    add_bullet(doc, b)

add_heading(doc, "Why It Was Done", level=2, color=DARK, size=12, space_before=8)
add_body(doc, (
    "External dependencies (S3, OpenAI, Snowflake) are inherently unreliable — they return "
    "transient errors, throttle bursts, or undergo brief outages. Without retry logic and "
    "proper connection handling, a single network hiccup could fail an entire website build "
    "and leave the user with no feedback. The /health endpoint is a prerequisite for "
    "container orchestrators (Kubernetes, ECS) and load-balancer health probes to route "
    "traffic only to live instances."
))

add_heading(doc, "Benefits", level=2, color=DARK, size=12, space_before=8)
for b in [
    "Website builds no longer fail silently due to a momentary S3 or OpenAI blip.",
    "Database migrations are reproducible and safe to run multiple times without side-effects.",
    "Ops teams can monitor liveness and readiness through a standard /health probe.",
    "Structured error logs dramatically reduce mean-time-to-diagnosis for production incidents.",
    "Graceful shutdown prevents in-flight requests from being dropped during rolling deploys.",
]:
    add_bullet(doc, b)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PRIORITY 3 — CODE QUALITY
# ══════════════════════════════════════════════════════════════════════════════
add_colored_heading(doc, "Priority 3 — Code Quality & Testing  (commit d3a0aba)", bg_hex="7C3AED")

add_heading(doc, "What Was Done", level=2, color=DARK, size=12, space_before=10)
for b in [
    "Created a pytest test suite (tests/) with 22 passing tests covering auth flows and website CRUD.",
    "Added a ResizeObserver smoke test (Node.js / jsdom) to validate the iframe layout observer.",
    "Extracted the 3 600-line inline <script> block from dashboard.html into frontend/dashboard.js.",
    "Added an error boundary around applySecEdits() to prevent one bad section edit from crashing the whole editor.",
    "Fixed a font-family name inconsistency that caused 'Inter' to render as the browser default.",
    "Enforced consistent code style with an .editorconfig and linting baseline.",
]:
    add_bullet(doc, b)

add_heading(doc, "Why It Was Done", level=2, color=DARK, size=12, space_before=8)
add_body(doc, (
    "A codebase without tests is a liability — every new feature or refactor risks silently "
    "breaking existing behaviour. Inline scripts block browser parsing and make code impossible "
    "to lint, minify, or cache independently. The font inconsistency caused visual regressions "
    "for end-users. Addressing these concerns before adding features ensures that the foundation "
    "is solid and that future contributors can verify correctness automatically."
))

add_heading(doc, "Benefits", level=2, color=DARK, size=12, space_before=8)
for b in [
    "22 automated tests catch regressions on every push without manual QA effort.",
    "Extracted dashboard.js enables tree-shaking, caching (content hash), and linting in CI.",
    "Error boundary in the section editor means one malformed section cannot break the entire page.",
    "Consistent fonts improve perceived design quality and brand alignment.",
    "Clear separation of HTML structure from JS logic reduces cognitive load for new developers.",
]:
    add_bullet(doc, b)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# PRIORITY 4 — PERFORMANCE & UX
# ══════════════════════════════════════════════════════════════════════════════
add_colored_heading(doc, "Priority 4 — Performance & UX Polish  (commit 153b550)", bg_hex="B45309")

add_heading(doc, "What Was Done", level=2, color=DARK, size=12, space_before=10)

sub_items = {
    "Item 16 — Lazy Section Editor Fields": [
        "openSecEditor() now shows an instant shimmer placeholder while the editor panel opens.",
        "All field HTML is collected into a string array (htmlChunks[]) and written with a single "
        "innerHTML assignment inside a requestAnimationFrame callback.",
        "Eliminates the previous O(n²) pattern where repeated innerHTML += caused the browser "
        "to re-parse and re-render all existing nodes on each append.",
    ],
    "Item 17 — Debounced Live Preview": [
        "Added a _debounce(fn, 120) utility function near the top of dashboard.js.",
        "oninput handlers for previewFieldStyle() and previewBg() are wrapped in debounced versions.",
        "Prevents the iframe from re-rendering on every keystroke; fires at most once per 120 ms.",
    ],
    "Item 18 — Paginated API Endpoints": [
        "list_clients, list_websites, and list_my_websites now accept page and limit query params.",
        "All three return a {items, total, page, pages} envelope for future infinite-scroll support.",
        "_fetchMyWebsites() helper in dashboard.js transparently unwraps the envelope.",
        "Default limit of 200 prevents unbounded result sets from large accounts.",
    ],
    "Item 19 — Server-Sent Events Build Stream": [
        "New GET /{website_id}/build-stream endpoint streams JSON status events every 2 seconds.",
        "Auth token passed as a ?token= query parameter (EventSource API cannot set headers).",
        "Both build-polling setInterval loops in dashboard.js replaced with EventSource.",
        "Auto-closes after 'built' or 'error' status, or after 5-minute timeout (150 ticks).",
    ],
    "Item 20 — Consolidated Toast Notifications": [
        "Shared toast() helper extracted to frontend/toast.js.",
        "Loaded as a standalone <script> before dashboard.js so all pages can reuse it.",
        "Prevents toast stacking via a clearTimeout guard on a per-element timer.",
    ],
}

for sub_title, sub_bullets in sub_items.items():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(sub_title)
    run.bold = True
    run.underline = True
    run.font.color.rgb = DARK
    run.font.size = Pt(10.5)
    for b in sub_bullets:
        add_bullet(doc, b)

add_heading(doc, "Why It Was Done", level=2, color=DARK, size=12, space_before=8)
add_body(doc, (
    "Performance and feedback quality directly impact user trust. A section editor that "
    "freezes for half a second on every click, or a build modal that requires manual refresh "
    "to see progress, creates friction that erodes confidence. Pagination prevents memory "
    "exhaustion and slow page loads as data grows. Debouncing protects the browser's "
    "rendering pipeline from thrashing during fast input. Consolidating toast avoids "
    "duplicate definitions across files drifting out of sync."
))

add_heading(doc, "Benefits", level=2, color=DARK, size=12, space_before=8)
for b in [
    "Section editor opens instantly with a shimmer; fields appear on the very next frame.",
    "Live style preview no longer saturates the CPU during fast typing.",
    "Build progress is real-time and automatic — no manual refresh required.",
    "API responses remain fast regardless of how many websites or clients a user accumulates.",
    "Single source of truth for toast notifications; visual consistency across all pages.",
]:
    add_bullet(doc, b)

add_divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Conclusion", level=1, color=INDIGO, size=18, space_before=14)
add_body(doc, (
    "The four priority phases represent a deliberate, systematic investment in product "
    "quality before expanding feature scope. Security was addressed first because "
    "vulnerabilities compound — a feature built on an insecure foundation inherits "
    "those vulnerabilities. Stability came next to ensure reliable customer-facing "
    "behaviour under real-world network conditions. Code quality enabled safe, fast "
    "iteration by providing automated regression coverage. Finally, performance and UX "
    "polish was applied to raise the standard of day-to-day interaction for builders "
    "using the platform."
))
add_body(doc, (
    "With all four priorities complete, the codebase is now well-positioned for the "
    "next phase of feature development: richer AI-driven content generation, multi-page "
    "site management, e-commerce integrations, and team collaboration — all built "
    "on a secure, stable, tested, and performant foundation."
))

doc.add_paragraph()

# Footer note
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    f"Agentic S3 Website Builder  ·  Engineering Report  ·  "
    f"{datetime.date.today().strftime('%B %Y')}"
)
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0xA0, 0xAE, 0xC0)
fr.italic = True

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = "docs/Engineering_Improvements_Report.docx"
import os; os.makedirs("docs", exist_ok=True)
doc.save(out_path)
print(f"Saved: {out_path}")
