"""
Generate a structured Word document capturing the agent design discussions
for the Agentic S3 Website Builder project.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1E, 0x29, 0x3B)   # headings
C_ACCENT = RGBColor(0x4F, 0x46, 0xE5)   # indigo accent
C_MUTED  = RGBColor(0x6B, 0x72, 0x80)   # muted grey
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_RED    = RGBColor(0xDC, 0x26, 0x26)
C_AMBER  = RGBColor(0xD9, 0x77, 0x06)
C_GREEN  = RGBColor(0x05, 0x96, 0x69)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(style=style_map.get(level, "Heading 1"))
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    run.font.color.rgb = color or C_DARK
    return p

def add_body(doc, text, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.bold = True
        br.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.bold = True
        br.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_divider(doc):
    p = doc.add_paragraph("─" * 90)
    run = p.runs[0]
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def add_table(doc, headers, rows, col_widths=None, header_bg="1E293B"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg   = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            p    = cell.paragraphs[0]
            # Support inline bold via ** markers
            parts = cell_text.split("**")
            for idx, part in enumerate(parts):
                r = p.add_run(part)
                r.bold = (idx % 2 == 1)
                r.font.size = Pt(9.5)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_info_box(doc, title, lines, bg="EEF2FF", border_color="4F46E5"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    title_run = p.add_run(f"  {title}\n")
    title_run.bold = True
    title_run.font.color.rgb = C_ACCENT
    title_run.font.size = Pt(11)
    for line in lines:
        lr = p.add_run(f"  {line}\n")
        lr.font.size = Pt(10)
        lr.font.color.rgb = C_DARK


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Agentic AI Website Builder")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = C_DARK

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("Agent Architecture — Design & Strategy Reference")
sr.font.size = Pt(14)
sr.font.color.rgb = C_ACCENT
sr.bold = True

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta_p.add_run(
    f"Project: senthilvasansubbu/agentic-s3-website-builder\n"
    f"Branch: main     |     Date: {datetime.date.today().strftime('%d %B %Y')}\n"
    f"Prepared from: AI Architecture Review Sessions"
)
mr.font.size = Pt(10)
mr.font.color.rgb = C_MUTED

doc.add_paragraph()
doc.add_paragraph()
add_divider(doc)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Table of Contents", 1, C_DARK)
toc_items = [
    ("1", "Current Agent Architecture", "3"),
    ("2", "Recommended Custom Agents — Full Catalogue", "4"),
    ("3", "Agent Execution Model", "6"),
    ("4", "Impact Analysis on the Current Application", "8"),
    ("5", "Infrastructure Pivot Resilience", "10"),
    ("6", "Strategic Recommendation — When to Go Agentic", "11"),
    ("7", "Incremental Implementation Roadmap", "12"),
]
toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.style = "Table Grid"
for ri, (num, title, pg) in enumerate(toc_items):
    row = toc_table.rows[ri]
    bg  = "F1F5F9" if ri % 2 == 0 else "FFFFFF"
    for ci in range(3):
        set_cell_bg(row.cells[ci], bg)
    row.cells[0].paragraphs[0].add_run(num).font.size  = Pt(10)
    row.cells[1].paragraphs[0].add_run(title).font.size = Pt(10)
    row.cells[2].paragraphs[0].add_run(pg).font.size   = Pt(10)
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(4.5)
    row.cells[2].width = Inches(0.6)
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — CURRENT STATE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Current Agent Architecture", 1, C_DARK)
add_body(doc,
    "The application currently uses a two-agent CrewAI pipeline defined in agents/crew.py. "
    "Both agents are instantiated fresh per build request and communicate sequentially — the designer's "
    "output becomes the developer's input context. Neither agent maintains state between requests.",
    size=10.5)
doc.add_paragraph()

add_heading(doc, "1.1 Existing Agents", 2, C_ACCENT)
add_table(doc,
    ["Agent", "File", "Role", "Task"],
    [
        ["**designer_agent**", "agents/designer_agent.py",
         "UI/UX Designer & Brand Content Strategist",
         "Produces detailed design spec: layout, colour palette, typography, image plan, content plan per section"],
        ["**developer_agent**", "agents/developer_agent.py",
         "Senior Full-Stack Web Developer & Content Engineer",
         "Converts design spec into a complete single-file HTML/CSS/JS website with real content, Unsplash images, booking form, contact section"],
    ],
    col_widths=[1.3, 1.6, 1.8, 2.8]
)
doc.add_paragraph()

add_heading(doc, "1.2 Current Build Flow", 2, C_ACCENT)
flow_steps = [
    ("HTTP POST /api/v1/websites/{id}/build", "Client submits build request with requirements, categories, location, social links"),
    ("Prompt assembly (website_builder.py)", "160 lines of inline Python build the LLM prompt — web search, social search, cart features, enrichment"),
    ("build_website(full_prompt)", "Synchronous blocking call to crew.py — freezes the HTTP request for 30–120 seconds"),
    ("Crew.kickoff()", "CrewAI runs designer_agent → developer_agent sequentially (Process.sequential)"),
    ("html_generator.py", "Writes output HTML to output/<project-slug>/index.html on disk"),
    ("HTTP response", "Returns { message, output_path, trace_id } when fully complete"),
]
add_table(doc,
    ["Step", "What Happens"],
    flow_steps,
    col_widths=[2.2, 5.3]
)
doc.add_paragraph()

add_heading(doc, "1.3 Current Limitations", 2, C_RED)
limitations = [
    "BackgroundTasks is imported in the route but never used — the build is fully blocking",
    "A single 60-second build can starve all other API requests on a single-worker uvicorn instance",
    "No job_id, no status polling endpoint — client has no way to track progress",
    "No QA validation — the generated HTML is served as-is with no correctness checks",
    "No SEO layer — meta tags, Open Graph, and schema.org markup are absent",
    "Prompt assembly logic is 160+ lines inline in the route handler — brittle and hard to test",
]
for l in limitations:
    add_bullet(doc, l)
doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — RECOMMENDED AGENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Recommended Custom Agents — Full Catalogue", 1, C_DARK)
add_body(doc,
    "The following agents were identified through analysis of the full application codebase — "
    "routes, services, tools, and frontend. They are grouped by pipeline stage.",
    size=10.5)
doc.add_paragraph()

add_heading(doc, "2.1 Website Build Pipeline Agents", 2, C_ACCENT)
build_agents = [
    ["**RequirementsAnalystAgent**", "Entry point",
     "Parses natural-language prompt + structured fields (categories, location, email, social links, scraped URL) into a clean, validated build brief. Replaces the 160-line inline prompt assembly in website_builder.py.",
     "Low — extracts existing logic"],
    ["**DesignerAgent**", "Step 2 (exists)",
     "UI/UX design specification: layout, colour palette, typography, image plan, content plan per section.",
     "None — already exists"],
    ["**MultiPageArchitectAgent**", "Step 3",
     "Plans multi-page site structure when num_pages > 1 or features like /blog, /livestream are enabled. Delegates each page to the DeveloperAgent.",
     "Medium — needs page routing logic"],
    ["**DeveloperAgent**", "Step 4 (exists)",
     "Converts design spec into complete HTML/CSS/JS. Currently outputs single-file; will output per-page files with MultiPageArchitect.",
     "Low — prompt extension only"],
    ["**SEOContentAgent**", "Step 5",
     "Generates SEO-optimised titles, meta descriptions, Open Graph tags, schema.org JSON-LD, and keyword-rich copy. Currently zero SEO coverage.",
     "Low — additive post-process"],
    ["**AssetOptimisationAgent**", "Step 6",
     "Post-build: compress images, inline critical CSS, generate sitemap.xml, robots.txt, manifest.json for PWA support.",
     "Medium — file I/O operations"],
    ["**QAReviewerAgent**", "Step 7",
     "Validates that all required sections exist (hero, categories, contact, form, footer), checks for broken Unsplash URLs, scores accessibility. Gates the build before deployment.",
     "Low — HTML parsing only"],
    ["**S3DeploymentAgent**", "Step 8",
     "Handles S3 bucket creation, file upload, CloudFront invalidation, DNS config, and deployment status. Calls notification_service.py on completion.",
     "Low — wraps existing s3_uploader.py"],
]
add_table(doc,
    ["Agent", "Stage", "Responsibility", "Change Impact"],
    build_agents,
    col_widths=[1.5, 1.0, 3.6, 1.4]
)
doc.add_paragraph()

add_heading(doc, "2.2 Shopping Cart Pipeline Agents", 2, C_ACCENT)
cart_agents = [
    ["**ProductCatalogAgent**", "POST /shop/import-catalog",
     "Wraps catalog_scraper.py — intelligently scrapes competitor/supplier URLs, normalises product names, prices, currencies, deduplicates, auto-assigns categories. Adds AI enrichment on top of the existing mechanical import.",
     "Low — wraps existing tool"],
    ["**PricingStrategyAgent**", "Post-catalog import",
     "Suggests compare_price, discount_pct, and flash_offer candidates based on margin rules or competitor pricing signals.",
     "Low — additive enrichment"],
    ["**CartUIGeneratorAgent**", "Website build (cart enabled)",
     "Specialised code generator for the storefront layer — cart drawer, product grid, filter sidebar, coupon input, checkout summary, order confirmation flow. Currently injected as a feature-flag string into the general DeveloperAgent prompt.",
     "Medium — separates cart UI from main build"],
    ["**CheckoutFlowAgent**", "Cart checkout",
     "Handles order summary → coupon validation → Stripe payment form → order reference → email/SMS/WhatsApp notification. Coordinates payment_service.py and notification_service.py.",
     "Medium — cross-service coordination"],
]
add_table(doc,
    ["Agent", "Trigger", "Responsibility", "Change Impact"],
    cart_agents,
    col_widths=[1.6, 1.4, 3.4, 1.1]
)
doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — EXECUTION MODEL
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Agent Execution Model", 1, C_DARK)

add_heading(doc, "3.1 Always-On Services vs. On-Demand Agents", 2, C_ACCENT)
add_body(doc,
    "Agents in this application are NOT microservices. They are short-lived Python objects "
    "instantiated inside a background worker, execute one task, then are garbage-collected. "
    "No agent keeps running between requests.",
    size=10.5)
doc.add_paragraph()

add_table(doc,
    ["Component", "Type", "Lifetime", "Examples"],
    [
        ["FastAPI app", "Always-on service", "Permanent", "app.py, all routes"],
        ["Celery / BackgroundTasks worker", "Always-on service", "Permanent", "Job queue processor"],
        ["Redis / DB", "Always-on service", "Permanent", "State, job status"],
        ["CrewAI Agents", "On-demand", "Per build job (~30–120s)", "designer_agent, developer_agent"],
        ["LLM calls (OpenAI)", "On-demand", "Per agent task (~5–60s)", "Each Task() in crew.py"],
        ["Tools (scraper, S3 upload)", "On-demand", "Per invocation (~1–10s)", "catalog_scraper, s3_uploader"],
    ],
    col_widths=[1.5, 1.3, 1.6, 3.1]
)
doc.add_paragraph()

add_heading(doc, "3.2 Recommended Async Build Architecture", 2, C_ACCENT)
add_body(doc,
    "The correct execution model replaces the current blocking HTTP call with a job queue pattern:",
    size=10.5)
doc.add_paragraph()

queue_steps = [
    ("1", "Client sends POST /build", "FastAPI validates request, saves status='queued' to DB"),
    ("2", "Job enqueued", "Build job pushed to Celery queue (or FastAPI BackgroundTasks)"),
    ("3", "Immediate HTTP response", "Returns { job_id, status: 'queued' } — no waiting"),
    ("4", "Worker picks up job", "Background worker process begins agent pipeline"),
    ("5", "Agent pipeline runs", "RequirementsAnalyst → Designer → Developer → SEO → QA → Deploy"),
    ("6", "DB updated", "status='published', live_url written to websites table"),
    ("7", "Notification sent", "notification_service.py emails/SMSes the user with live URL"),
    ("8", "Client polls status", "GET /websites/{id}/build-status returns current stage + result"),
]
add_table(doc,
    ["Step", "Action", "Detail"],
    queue_steps,
    col_widths=[0.4, 1.8, 5.3]
)
doc.add_paragraph()

add_heading(doc, "3.3 Event-Trigger Map", 2, C_ACCENT)
trigger_map = [
    ["POST /websites/{id}/build",        "RequirementsAnalyst → Designer → MultiPageArchitect → Developer → SEO → AssetOptimiser → QA → S3Deploy"],
    ["POST /shop/import-catalog",        "ProductCatalogAgent → PricingStrategyAgent"],
    ["POST /shop/cart-items (cart=true)", "CartUIGeneratorAgent → CheckoutFlowAgent"],
    ["PUT /websites/{id} (content edit)", "SEOContentAgent only"],
    ["Build complete + hosting_env=s3",  "S3DeploymentAgent only"],
]
add_table(doc,
    ["Trigger Event", "Agents Invoked"],
    trigger_map,
    col_widths=[2.2, 5.3]
)
doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — IMPACT ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Impact Analysis on the Current Application", 1, C_DARK)

add_heading(doc, "4.1 Immediate Problems (Existing Code)", 2, C_RED)
problems = [
    ["Blocking HTTP build call",
     "api/routes/website_builder.py L369",
     "build_website() called directly in async handler. BackgroundTasks imported but unused.",
     "🔴 HIGH — affects all users during any active build"],
    ["No build status feedback",
     "website_builder.py, dashboard.html",
     "No job_id, no polling endpoint, no WebSocket push. Client just waits.",
     "🔴 HIGH — poor UX, browser timeouts on slow builds"],
    ["CORS wildcard",
     "app.py L83",
     "allow_origins=['*'] — any domain can call the API.",
     "🟡 MEDIUM — security risk in production"],
    ["160-line inline prompt assembly",
     "website_builder.py L195–360",
     "Cart features, blog, chatbot, social links all concatenated inline. Breaks with every new feature.",
     "🟡 MEDIUM — maintainability risk"],
    ["No schema migration versioning",
     "database/migrations.py",
     "Schema changes are applied as raw DDL with no version tracking.",
     "🟡 MEDIUM — data loss risk on repeated runs"],
]
add_table(doc,
    ["Problem", "Location", "Description", "Risk"],
    problems,
    col_widths=[1.5, 1.6, 3.0, 1.4]
)
doc.add_paragraph()

add_heading(doc, "4.2 File-by-File Agent Integration Impact", 2, C_ACCENT)
file_impact = [
    ["agents/crew.py", "🔴 Major",
     "Single Crew([designer, developer]) expands to 8–10 agent pipeline. _generate_static_fallback() (200+ lines) stays untouched."],
    ["api/routes/website_builder.py", "🔴 Major",
     "build route must save status='queued', push to background, return job_id. New GET /build-status endpoint needed."],
    ["database/migrations.py", "🟡 Medium",
     "New columns needed: build_status, build_job_id, build_started_at, build_completed_at, build_error."],
    ["frontend/dashboard.html", "🟡 Medium",
     "JS must replace blocking fetch with polling/WebSocket. Add progress bar showing current agent stage."],
    ["app.py", "🟢 Low",
     "Additive only. If Celery added, register startup event. No breaking changes."],
    ["config/settings.py", "🟢 Low",
     "New env vars: REDIS_URL, MAX_CONCURRENT_BUILDS, BUILD_TIMEOUT_SECONDS. Additive."],
    ["services/", "🟢 None",
     "Agents call existing services (notification, hosting, image). No changes to service files."],
    ["agents/designer_agent.py", "🟢 None",
     "No changes. Stays as-is and slots into new pipeline."],
    ["agents/developer_agent.py", "🟢 None",
     "No changes unless multi-page output is required (prompt extension only)."],
]
add_table(doc,
    ["File / Area", "Risk Level", "Change Required"],
    file_impact,
    col_widths=[1.9, 0.9, 4.7]
)
doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — INFRASTRUCTURE PIVOT RESILIENCE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Infrastructure Pivot Resilience", 1, C_DARK)
add_body(doc,
    "A key architectural strength of this codebase is that agents produce HTML strings — "
    "they are completely decoupled from the hosting layer. Any infrastructure pivot affects "
    "only hosting_service.py and the uploader tool, never the agents themselves.",
    size=10.5)
doc.add_paragraph()

pivot_table = [
    ["S3 → Azure Blob Storage", "**Zero**",
     "Replace s3_uploader.py with azure_uploader.py. Update hosting_service.py. Agents untouched."],
    ["S3 → Google Drive", "**Zero**",
     "New gdrive_uploader.py. DNS routing config change. Agents untouched."],
    ["File system → Store HTML in DB + DNS by website_id", "**Zero**",
     "hosting_service.py writes to DB instead of disk. URL-router middleware in app.py. Agents untouched. This is the most future-proof architecture."],
    ["Single-file HTML → Multi-page output", "**Low**",
     "html_generator.py, crew.py return shape change. Developer agent prompt extension. Agents mostly untouched."],
    ["Plain HTML → React/Vue output", "**High**",
     "Both agent prompts completely rewritten. Output pipeline changes significantly. Build toolchain (npm, bundler) required."],
]
add_table(doc,
    ["Infrastructure Pivot", "Agent Impact", "What Actually Changes"],
    pivot_table,
    col_widths=[2.0, 0.8, 4.7]
)
doc.add_paragraph()

add_info_box(doc, "Key Insight",
    ["The agents are insulated from infrastructure because they produce content, not deployment.",
     "The hosting layer is a thin wrapper around the agent output — change the wrapper, not the agents.",
     "The 'store in DB + DNS by website_id' pattern is recommended as the most cloud-agnostic approach."])
doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — STRATEGIC RECOMMENDATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Strategic Recommendation — When to Go Agentic", 1, C_DARK)

add_heading(doc, "6.1 The Core Question", 2, C_ACCENT)
add_body(doc,
    "The application is currently at approximately 20–30% of its intended feature set. "
    "The question is whether to invest in a full multi-agent pipeline now or after the application matures.",
    size=10.5)
doc.add_paragraph()

add_heading(doc, "6.2 Recommendation: Build Agents After Application Matures", 2, C_GREEN)
reasons = [
    ("Application contract is still changing",
     "Every new feature (livestream, blog, chatbot, cart features) adds to the agent prompt. "
     "Adding more agents now means updating N agent task descriptions with every change instead of 1."),
    ("Cost of agentic complexity is asymmetric",
     "Current: 2 LLM calls per build (~$0.10–0.40). A 10-agent pipeline = 10 LLM calls. "
     "At a development stage where prompts change frequently, you pay to refine 10 prompts instead of 2."),
    ("No production data to guide agent boundaries",
     "The best agent designs come from real failure patterns. Without live user traffic, "
     "you're building agents to solve problems you haven't confirmed exist yet."),
    ("Infrastructure isn't ready to support longer pipelines",
     "The current blocking HTTP call means even a 2-agent build risks timeout. "
     "A 10-agent pipeline is unmaintainable without the background job queue in place first."),
]
for title, detail in reasons:
    add_bullet(doc, f" {detail}", bold_prefix=f"{title}: ")
doc.add_paragraph()

add_heading(doc, "6.3 The One Exception — Do This Now", 2, C_AMBER)
add_body(doc,
    "The only agent worth adding today is a thin RequirementsAnalystAgent — not because of AI "
    "complexity, but because the prompt-building logic in website_builder.py (lines 195–360) is "
    "160 lines of inline Python string manipulation that breaks with every new feature. "
    "Extracting it into a structured agent class makes the code maintainable regardless of feature changes.",
    size=10.5)
doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Incremental Implementation Roadmap", 1, C_DARK)

add_heading(doc, "Phase 1 — Fix the Foundation (Now, 1–2 days)", 2, C_ACCENT)
add_body(doc, "Zero breaking changes. Fixes the most critical production risk.", italic=True, color=C_MUTED)
p1 = [
    "website_builder.py: wrap build_website() in background_tasks.add_task() — use the already-imported BackgroundTasks that is currently unused",
    "Add GET /websites/{id}/build-status polling endpoint",
    "DB: add build_status column (queued | running | built | error) with default 'idle'",
    "dashboard.html: replace blocking fetch with polling loop + progress spinner",
    "app.py: tighten CORS allow_origins from ['*'] to specific domains",
]
for s in p1:
    add_bullet(doc, s)
doc.add_paragraph()

add_heading(doc, "Phase 2 — Stabilise Features (Now → 60% complete)", 2, C_ACCENT)
add_body(doc, "Do not add agents yet. Focus on feature completeness.", italic=True, color=C_MUTED)
p2 = [
    "Lock the feature set: cart, blog, chatbot, livestream, multi-page",
    "Stabilise the DB schema — add proper migration versioning",
    "Get real users building real websites on the platform",
    "Collect logs of where build quality fails (missing sections, broken images, bad SEO)",
    "Extract RequirementsAnalystAgent from the inline prompt assembly logic only",
]
for s in p2:
    add_bullet(doc, s)
doc.add_paragraph()

add_heading(doc, "Phase 3 — Introduce Agents (60–70% complete)", 2, C_ACCENT)
add_body(doc, "Application is stable. Features locked. Failure patterns known.", italic=True, color=C_MUTED)
p3 = [
    "Add RequirementsAnalystAgent (if not done in Phase 2)",
    "Add QAReviewerAgent — you now know what 'bad output' looks like from real builds",
    "Add SEOContentAgent — you have confirmed SEO gaps from real traffic data",
    "Wire new agents into crew.py pipeline without changing existing agent prompts",
]
for s in p3:
    add_bullet(doc, s)
doc.add_paragraph()

add_heading(doc, "Phase 4 — Full Agentic Pipeline (80–90% complete)", 2, C_ACCENT)
add_body(doc, "Production-ready, real traffic, real cart data, infrastructure decided.", italic=True, color=C_MUTED)
p4 = [
    "Add S3DeploymentAgent (or Azure/DB equivalent based on final infrastructure decision)",
    "Add ProductCatalogAgent — real catalog imports, real edge cases now known",
    "Add CartUIGeneratorAgent — cart feature set locked, real checkout failures observed",
    "Add CheckoutFlowAgent — payment integration stable",
    "Add MultiPageArchitectAgent — multi-page structure confirmed and tested",
    "Add AssetOptimisationAgent — performance bottlenecks identified from real usage",
]
for s in p4:
    add_bullet(doc, s)
doc.add_paragraph()

add_heading(doc, "Phase 4 — Maturity Threshold Summary", 2, C_DARK)
roadmap_summary = [
    ["Now (20–30%)",    "Fix blocking build call, stabilise features, lock schema, real users"],
    ["60–70% complete", "RequirementsAnalyst, QAReviewer, SEOContent agents"],
    ["80–90% complete", "S3Deploy, ProductCatalog, CartUI, Checkout, MultiPage, AssetOptimisation agents"],
    ["100% (full SaaS)", "Full 10-agent pipeline, Celery workers, WebSocket progress, real-time monitoring"],
]
add_table(doc,
    ["Application Maturity", "Agent Actions at This Stage"],
    roadmap_summary,
    col_widths=[1.8, 5.7]
)
doc.add_paragraph()

add_divider(doc)
add_body(doc,
    "Core principle: Add agents when the application tells you where it hurts — not before. "
    "Ship features, get users, let real failure patterns guide the agent design.",
    italic=True, color=C_MUTED, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "docs/Agent_Architecture_Reference.docx"
import os; os.makedirs("docs", exist_ok=True)
doc.save(out_path)
print(f"✅  Saved: {out_path}")
