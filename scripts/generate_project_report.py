"""
Generate a structured Word document summarising the project discussion,
technical decisions, completed work, and the security + improvement plan.
Run: python3 scripts/generate_project_report.py
Output: docs/project_discussion_report.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

os.makedirs("docs", exist_ok=True)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:  run.font.color.rgb = RGBColor(*color)
    if size:   run.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def table_2col(rows, header=None):
    cols = len(header) if header else len(rows[0])
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    t.style = "Table Grid"
    if header:
        for i, h in enumerate(header):
            cell = t.rows[0].cells[i]
            cell.text = h
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            t.rows[r_idx + (1 if header else 0)].cells[c_idx].text = str(val)
    doc.add_paragraph()

def divider():
    doc.add_paragraph("─" * 80)

# ═════════════════════════════════════════════════════════════════════════════
# COVER
# ═════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("Agentic S3 Website Builder")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.add_run("Project Discussion, Technical Decisions & Improvement Plan").italic = True

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
heading("1. Project Overview")
para(
    "The Agentic S3 Website Builder is a FastAPI-based SaaS platform that uses AI agents "
    "(CrewAI + OpenAI GPT-4) to autonomously build, style, and deploy websites for clients. "
    "It includes a visual dashboard editor, multi-tenant authentication, e-commerce, payments, "
    "monitoring, and AWS S3 deployment."
)

heading("1.1 Technology Stack", level=2)
table_2col([
    ("Backend",        "Python 3, FastAPI, Uvicorn"),
    ("AI / Agents",    "CrewAI 1.14, OpenAI GPT-4 / gpt-4o-mini"),
    ("Frontend",       "Vanilla JS, single dashboard.html (~5000 lines), iframe preview editor"),
    ("Database",       "Snowflake (prod) / SQLite auto-fallback (dev)"),
    ("Storage",        "Local uploads (dev) / AWS S3 (prod target)"),
    ("Payments",       "Stripe subscriptions + storefront checkout"),
    ("Notifications",  "Twilio (SMS / WhatsApp) + SMTP email"),
    ("Auth",           "JWT (PyJWT) + OTP (email/SMS), bcrypt password hashing"),
    ("Testing",        "Node.js + JSDOM — 8 tests on editor logic"),
    ("Deployment",     "AWS S3 static hosting + CloudFront (planned)"),
], header=["Layer", "Technology"])

heading("1.2 User Role Hierarchy", level=2)
table_2col([
    ("superuser",  "Platform administrator — manages app_users, has console access"),
    ("app_user",   "Agency / SaaS operator — builds websites, manages billing, onboards clients"),
    ("client",     "End-business owner — can edit their own website content"),
    ("customer",   "Shopper on a client's storefront — no dashboard access"),
], header=["Role", "Description"])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 2. DEVELOPMENT SESSION SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("2. Development Session Summary (Recent Conversations)")

sessions = [
    (
        "Session 1 — Dashboard Editor Bug Fixes",
        [
            "Fixed applySecEdits(): fieldNum was not advancing past link fields, causing image field IDs to be misaligned. Fix: added fieldNum += linkDivs.length before the images loop.",
            "Fixed style bar not reflecting changes: template CSS uses !important which overrides plain el.style.prop. Fix: changed all style bar assignments to setProperty('prop', value, 'important').",
            "Fixed color picker live preview: added oninput alongside onchange so color changes show in real time.",
            "Fixed brand name style in the nav section (section 0) editor.",
        ]
    ),
    (
        "Session 2 — Image Editor Modal",
        [
            "Built a full image editor modal with three tabs: Crop (Cropper.js 1.6.2), Resize, and Enhance (brightness/contrast/saturation).",
            "Fixed image editor callback bug: _imgEdClose() was nulling _imgCallback before the callback was invoked. Fix: saved callback reference to a local variable before closing.",
            "Intercepted uploadSectionImage() and uploadBgImage() to open the image editor first, then upload the processed result.",
            "Password reset performed for sayeesaran.s@gmail.com → Sayee@1234.",
            "Git commit and push completed (commit a92679c).",
        ]
    ),
    (
        "Session 3 — Responsive Design & Hamburger Menu",
        [
            "_injectResponsiveEnhancements(doc) function added to dashboard: injected into iframe on every load to add responsive CSS and replace the hamburger onclick handler.",
            "Added full tablet (900px) and mobile (640px) breakpoints for grids, hero, nav, footer.",
            "Fixed hamburger white space bug: .nav-links { display: none !important } added to the mobile CSS block so it is hidden by default and CSS takes over on close (removeProperty).",
            "Fixed hamburger collapse: added ResizeObserver on doc.body to show/hide hamburger based on actual iframe pixel width (not CSS media query, which doesn't work inside the dashboard iframe).",
            "Fixed close-on-link-click: used capture-phase click listener on navLinks + hashchange listener + outside-click dismiss on the document.",
            "Fixed Google Fonts URL: font_heading / font_body values were full CSS stacks ('Poppins', sans-serif) — stripped to just the font name before building the Google Fonts URL.",
            "Added node_modules/ to .gitignore (was being tracked).",
            "Git commit and push completed (commit ee8d8c8).",
        ]
    ),
]

for title, items in sessions:
    heading(title, level=2)
    for item in items:
        bullet(item)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3. ARCHITECTURAL DECISIONS
# ═════════════════════════════════════════════════════════════════════════════
heading("3. Key Architectural Decisions")

decisions = [
    ("Image Storage",
     "Decided to keep images in local uploads (data/uploads/) for now. Migration to S3 planned as Priority 2 item once security hardening is complete. Image service already has S3 upload logic built in."),
    ("Database",
     "Using SQLite locally via the auto-fallback in database/snowflake_client.py. The SQL adapter layer converts Snowflake-specific syntax (UUID_STRING, TIMESTAMP_NTZ, VARIANT) to SQLite automatically. Production will use Snowflake."),
    ("Test Environment",
     "Decision: use a 'dev' branch (not a separate repository) when the time comes. Recommended to set up once the first paying client is onboarded or payment processing is introduced. Current focus is main branch."),
    ("Frontend Architecture",
     "Remains vanilla JS in a single dashboard.html. No framework or bundler required at this stage. Splitting into partials (dashboard.js) is listed as a Priority 3 maintainability task."),
    ("Rate Limiting",
     "Not yet implemented. Agreed to add slowapi-based rate limiting on auth, OTP, and chatbot endpoints as the first security hardening step."),
]

for title, text in decisions:
    heading(title, level=2)
    para(text)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 4. SYSTEM LIMITS & CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════════
heading("4. System Limits & Configuration")

heading("4.1 Current API Limits", level=2)
table_2col([
    ("Image upload max size",     "10 MB  (services/image_service.py)"),
    ("Chatbot max_tokens",        "512  (api/routes/chatbot.py) — flagged as too tight"),
    ("Currency API timeout",      "4 seconds"),
    ("Monitoring check timeout",  "Configurable per check type"),
    ("Free plan page limit",      "10 pages  (api/routes/website_builder.py)"),
    ("Rate limiting",             "NONE — to be added (Priority 1)"),
], header=["Limit", "Value / Location"])

heading("4.2 Disk Usage (Dev Container)", level=2)
table_2col([
    ("Total disk",        "32 GB"),
    ("Used",              "12 GB  (38%)"),
    ("Available",         "19 GB"),
    ("Project size",      "~10 MB (excluding node_modules)"),
    ("Uploaded images",   "~8.5 MB in data/uploads/"),
], header=["Metric", "Value"])

heading("4.3 Required Environment Variables", level=2)
table_2col([
    ("OPENAI_API_KEY",        "OpenAI API key for website generation and chatbot"),
    ("AWS_ACCESS_KEY_ID",     "AWS credentials for S3 deployment"),
    ("AWS_SECRET_ACCESS_KEY", "AWS credentials for S3 deployment"),
    ("AWS_REGION",            "AWS region (default: us-east-1)"),
    ("S3_BUCKET_NAME",        "S3 bucket for website hosting"),
    ("WEBSITE_DOMAIN",        "Public domain for hosted sites"),
    ("SNOWFLAKE_ACCOUNT",     "Snowflake account ID (leave blank to use SQLite)"),
    ("SNOWFLAKE_USER",        "Snowflake username"),
    ("SNOWFLAKE_PASSWORD",    "Snowflake password"),
    ("STRIPE_SECRET_KEY",     "Stripe secret key for payments"),
    ("STRIPE_WEBHOOK_SECRET", "Stripe webhook signing secret"),
    ("TWILIO_ACCOUNT_SID",    "Twilio SID for SMS/WhatsApp OTP"),
    ("TWILIO_AUTH_TOKEN",     "Twilio auth token"),
    ("TWILIO_FROM_NUMBER",    "Twilio sender number"),
    ("CORS_ORIGINS",          "Comma-separated allowed origins (blank = localhost only)"),
], header=["Variable", "Purpose"])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 5. TODO PLAN
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Prioritised TODO Plan")

para(
    "Following a deep code analysis, the items below were identified and agreed upon. "
    "Security items are to be completed before any new feature work.",
    bold=False
)

heading("🔴 Priority 1 — Security (Before Any Real Users)", level=2)
p1 = [
    ("1", "Add rate limiting on /auth/login, /auth/register, /auth/verify-otp, /chatbot using slowapi (3 lines per route)"),
    ("2", "Fix feedback privacy — GET /feedback/{website_id} requires no auth; add website-ownership check"),
    ("3", "Fix monitoring superuser check — uses plan field check instead of role == 'superuser'"),
    ("4", "Add chatbot input length limit — cap message at 500 chars, sanitize before OpenAI call"),
    ("5", "Create .env.example file documenting all required environment variables"),
]
table_2col(p1, header=["#", "Task"])

heading("🟠 Priority 2 — Stability (Before Client Onboarding)", level=2)
p2 = [
    ("6",  "Move image uploads to S3 — image_service.py already has logic, change upload destination"),
    ("7",  "Add DB migration versioning — schema_version table + version-checked runner"),
    ("8",  "Add OpenAI retry logic in build_website() — exponential backoff, 3 attempts"),
    ("9",  "Increase chatbot max_tokens from 512 to 1024 or 2048"),
    ("10", "Add /health endpoint returning DB status, disk usage, and service versions"),
]
table_2col(p2, header=["#", "Task"])

heading("🟡 Priority 3 — Quality & Maintainability", level=2)
p3 = [
    ("11", "Python API test suite — pytest covering auth flow, website CRUD, payment webhook"),
    ("12", "Update hamburger JS test (Test 7) to cover ResizeObserver path"),
    ("13", "Extract dashboard JS into dashboard.js partial — reduce 5000-line file size"),
    ("14", "Add error boundary in applySecEdits() — silent JS failures give no user feedback"),
    ("15", "Fix font_heading/font_body at source in agents/crew.py — store as plain name not CSS stack"),
]
table_2col(p3, header=["#", "Task"])

heading("🟢 Priority 4 — Features (Growth Phase)", level=2)
p4 = [
    ("16", "Custom domain CNAME automation — hosting_service.py has instructions but no auto flow"),
    ("17", "Multi-page website support — all sites currently single-page index.html"),
    ("18", "Dashboard mobile layout — the dashboard UI itself needs responsive design"),
    ("19", "Scheduled monitoring alerts — APScheduler installed but not wired to email/SMS on downtime"),
    ("20", "Create 'dev' branch + GitHub Actions CI — trigger after Priority 1 & 2 complete"),
]
table_2col(p4, header=["#", "Task"])

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 6. DECISION LOG
# ═════════════════════════════════════════════════════════════════════════════
heading("6. Decision Log")

table_2col([
    ("Apr 21 2026", "Security before features",   "Do Priority 1-2 security items before any new feature work. Real client (MMGK) already using platform; feedback/auth endpoints are exposed."),
    ("Apr 21 2026", "Test environment timing",     "Set up dev branch after first paying client or before payment processing changes. Not needed immediately."),
    ("Apr 21 2026", "Image storage",               "Keep local for now. Migrate to S3 as Priority 2 item."),
    ("Apr 21 2026", "Branch vs separate repo",     "Use a branch (dev) not a separate repository. Same codebase, easier to merge."),
    ("Apr 21 2026", "node_modules tracking",       "Added node_modules/ to .gitignore — was being committed accidentally."),
], header=["Date", "Decision", "Notes"])

# ═════════════════════════════════════════════════════════════════════════════
# 7. NEXT STEPS
# ═════════════════════════════════════════════════════════════════════════════
heading("7. Immediate Next Steps")
bullet("Implement Priority 1 items (security) — estimated 1–2 days")
bullet("Item 1: slowapi rate limiting on auth + chatbot routes")
bullet("Item 2: feedback endpoint auth/ownership check")
bullet("Item 3: monitoring superuser role fix")
bullet("Item 4: chatbot input sanitization + length cap")
bullet("Item 5: create .env.example")
bullet("Then proceed to Priority 2 (stability) — estimated 2–3 days")
bullet("Then return to feature development")

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
out_path = "docs/project_discussion_report.docx"
doc.save(out_path)
print(f"✅  Saved: {out_path}")
